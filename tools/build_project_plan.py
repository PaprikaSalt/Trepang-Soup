from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Trepang Soup 项目规划.docx"

FONT = "Microsoft YaHei"
NAVY = "17242F"
BLUE = "29485C"
AMBER = "D7A95B"
INK = "22313A"
MUTED = "65737C"
PALE_BLUE = "E8EEF2"
PALE_AMBER = "FBF4E7"
PALE_GRAY = "F3F5F6"
WHITE = "FFFFFF"
BORDER = "CBD4D9"

# compact_reference_guide token map
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 80
CELL_MARGIN_BOTTOM_DXA = 80
CELL_MARGIN_SIDE_DXA = 120


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str = INK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    """Apply both Latin and CJK fonts so Word and LibreOffice render consistently."""
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_DXA),
        ("bottom", CELL_MARGIN_BOTTOM_DXA),
        ("start", CELL_MARGIN_SIDE_DXA),
        ("end", CELL_MARGIN_SIDE_DXA),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    """Lock Word table geometry to explicit DXA values, including every cell."""
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA} DXA")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("TREPANG SOUP  /  项目规划")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)


def add_title(doc: Document) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("PRODUCT & IMPLEMENTATION PLAN")
    set_run_font(run, size=10, color=AMBER, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Trepang Soup")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("海龟汤多人在线桌面平台")
    set_run_font(run, size=16, color=BLUE, bold=True)

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.left_indent = Inches(0.65)
    statement.paragraph_format.right_indent = Inches(0.65)
    statement.paragraph_format.space_after = Pt(66)
    run = statement.add_run("一个无需注册、为固定朋友群体设计的合作推理房间，由 DeepSeek 驱动 AI 主持人。")
    set_run_font(run, size=11.5, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(3)
    run = meta.add_run("版本 1.0  ·  2026 年 7 月 30 日")
    set_run_font(run, size=10.5, color=INK, bold=True)

    stage = doc.add_paragraph()
    stage.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stage.add_run("阶段：规划定稿 / Windows 客户端开发启动")
    set_run_font(run, size=9.5, color=MUTED, italic=True)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_callout(doc: Document, label: str, text: str, *, fill: str = PALE_AMBER) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=fill, size=2)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}  ")
    set_run_font(lead, size=10.5, color=NAVY, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=NAVY, bold=True)

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            if len(table.rows) % 2 == 0:
                set_cell_fill(cell, WHITE)
            else:
                set_cell_fill(cell, PALE_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index == 0 and len(headers) > 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=INK, bold=(index == 0 and len(headers) > 1))

    # Rows are appended after the initial geometry pass, so apply widths again.
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_section_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_document() -> Document:
    doc = Document()
    apply_document_styles(doc)
    add_running_furniture(doc.sections[0])
    add_title(doc)

    add_heading(doc, "1. 项目定义")
    add_callout(
        doc,
        "一句话目标",
        "为不超过 20 人的固定朋友群体提供一个打开即玩、无需注册、由 AI 主持的多人合作海龟汤桌面房间。",
    )
    add_body(
        doc,
        "Trepang Soup 不追求公开社区、用户增长或商业化。产品优先级依次是：便捷、稳定、氛围、美观、易维护。所有设计均以小规模好友房间为前提。",
    )

    add_heading(doc, "1.1 产品原则", level=2)
    add_bullets(
        doc,
        [
            "零账号负担：玩家只填写昵称和邀请码。",
            "合作而非竞争：所有线索、提示和结算均面向全房间公开。",
            "主持人克制：AI主要判断问题，不用长篇输出抢走玩家的推理过程。",
            "客户端优先：Windows 程序必须流畅、简洁并具有统一的视觉气质。",
            "有限规模优先：最多 20 人，不为未来扩容牺牲当前实现的清晰度。",
        ],
    )

    add_heading(doc, "1.2 首版明确不做", level=2)
    add_bullets(
        doc,
        [
            "注册、登录、账号找回和云端个人资料。",
            "网页端、移动端、应用商店发布和自动更新。",
            "语音聊天、语音输入、语音识别或 AI 语音播报。",
            "排行榜、长期积分、竞技机制和付费系统。",
            "对局记录云端保存及 Markdown、图片、PDF 导出。",
            "HTTPS/WSS 与 Docker 部署；协议会为以后升级保留边界。",
        ],
    )

    add_heading(doc, "2. 用户与权限")
    add_table(
        doc,
        ["角色", "权限与责任"],
        [
            ("普通玩家", "加入房间、参与讨论、提交或撤回待回答问题、请求提示、发起结案。"),
            ("房主", "拥有普通玩家能力，并可选择题目来源、开始游戏、移除玩家和关闭房间。"),
            ("管理员", "通过隐藏入口与专用密码管理私人题库；管理员身份不等于房主身份。"),
            ("AI主持人", "掌握汤底、依序回答问题、生成公共提示、判断结案并完成赛后总结。"),
        ],
        [2160, 7200],
    )

    add_heading(doc, "2.1 无账号身份机制", level=2)
    add_body(
        doc,
        "创建或加入房间时，客户端生成本机临时身份标识；服务端返回房间会话令牌。断线后，同一台电脑使用令牌恢复昵称和房间身份。昵称在单个房间内必须唯一。服务器重启后令牌与房间同时失效。",
    )

    add_heading(doc, "3. 完整游戏流程")
    add_numbered(
        doc,
        [
            "玩家启动程序并填写昵称。",
            "玩家创建房间，或填写 6 位易读邀请码加入好友房间。",
            "房主选择“AI 自动生成”或“私人题库随机抽取”。",
            "AI生成题目时，房主分别选择难度和风格；私人题库不显示筛选项。",
            "房主开始游戏，所有玩家看到汤面、主持人问答区、待回答队列和讨论区。",
            "玩家在讨论区自由交流，或把正式问题提交到主持人队列。",
            "服务端按接收时间串行送交 AI；等待中的问题可由提交者撤回。",
            "任意玩家可请求公共提示，AI梳理已确认事实并给出轻度提示。",
            "任意玩家可发起“我们知道真相了”，提交完整推理供 AI 判断。",
            "结案成功或放弃后公布汤底，AI生成评分、奖项和简短复盘。",
            "每台参与电脑在本地保存完整对局记录，服务端清除本局运行数据。",
        ],
    )

    add_heading(doc, "4. 房间与生命周期")
    add_table(
        doc,
        ["能力", "首版规则", "异常与边界"],
        [
            ("邀请码", "6 位字符，排除 0/O、1/I 等易混字符。", "邀请码只在活跃房间内有效。"),
            ("人数", "单房间最多 20 人。", "达到上限时拒绝新加入者并显示明确原因。"),
            ("中途加入", "游戏开始后仍可加入，并同步完整房间历史。", "同步完成前界面进入只读加载状态。"),
            ("断线重连", "同一电脑自动尝试恢复身份。", "超过服务端保留窗口或服务器重启后无法恢复。"),
            ("房主转移", "房主离线一段时间后转交给最早加入的在线玩家。", "原房主回来后作为普通玩家恢复。"),
            ("关闭房间", "房主可主动关闭。", "所有客户端收到关闭原因并返回首页。"),
        ],
        [1440, 3960, 3960],
    )

    add_heading(doc, "5. 游戏主界面")
    add_body(
        doc,
        "主界面使用三栏结构，并根据窗口宽度自适应。核心线索始终处于视觉中心，讨论消息不会挤占正式问答的阅读空间。",
    )
    add_table(
        doc,
        ["区域", "内容", "主要交互"],
        [
            ("主持人问答区", "汤面、已回答问题、AI回复、公共提示和系统事件。", "阅读、滚动、定位新回答。"),
            ("待回答队列", "按服务器接收顺序显示等待中与思考中的问题。", "提交者可撤回等待问题；思考中不可撤回。"),
            ("玩家讨论区", "玩家之间的实时文字聊天。", "发送讨论消息；不进入主持人问题队列。"),
        ],
        [1680, 4560, 3120],
    )

    add_heading(doc, "5.1 问题状态机", level=2)
    add_table(
        doc,
        ["状态", "含义", "允许操作"],
        [
            ("本地草稿", "尚未提交。", "编辑、清空、提交。"),
            ("等待回答", "服务端已接收，尚未送给 AI。", "提交者撤回。"),
            ("AI思考中", "已锁定并正在处理。", "任何人都不能撤回。"),
            ("已回答", "AI结果已广播。", "进入正式问答历史。"),
            ("失败", "调用失败或响应无效。", "系统自动有限重试，随后允许重新提交。"),
        ],
        [1680, 4560, 3120],
    )

    add_heading(doc, "6. AI主持人行为")
    add_callout(
        doc,
        "主持语气",
        "温和、有一点损友感，但不抢戏。可以轻微调侃绕圈推理，不根据昵称做人身化发挥，不使用冒犯性语言。",
        fill=PALE_BLUE,
    )
    add_bullets(
        doc,
        [
            "正式回答以“是 / 否 / 无关 / 部分正确”为核心，附加说明保持简短。",
            "AI只把正式问答视为已确认事实；玩家讨论不自动成为事实。",
            "提示会梳理已确认信息、指出矛盾或遗漏方向，但不直接公布关键真相。",
            "结案判断分为正确、接近和错误；接近时只指出缺失的关键拼图。",
            "模型输出必须经过服务端结构校验，不能由客户端直接信任自然语言结果。",
        ],
    )

    add_heading(doc, "6.1 “我没招了”公共提示", level=2)
    add_body(
        doc,
        "任意玩家均可无限次点击。提示进入主持人问答区并对全房间可见，系统记录请求者。每次使用会降低本局整体评分，但不产生个人惩罚或冷却时间。",
    )

    add_heading(doc, "6.2 赛后结算", level=2)
    add_bullets(
        doc,
        [
            "本局评分：结合有效问题、推理效率、提示使用次数和是否放弃。",
            "MVP玩家：对还原真相贡献最大者。",
            "最有价值问题：明显推进推理的问题。",
            "最佳带偏奖：有趣但把大家引向错误方向的问题或推理。",
            "小总结：回顾关键转折，语气轻松但不冒犯。",
        ],
    )

    add_heading(doc, "7. 内容系统")
    add_heading(doc, "7.1 AI自动生成", level=2)
    add_table(
        doc,
        ["维度", "选项"],
        [
            ("难度", "新手：线索集中；标准：需组合多条信息；烧脑：允许误导视角与多层因果。"),
            ("风格", "轻松日常、经典悬疑、暗黑惊悚、荒诞幽默。"),
            ("质量门槛", "先生成汤面、汤底与关键真相，再执行逻辑自洽检查；未通过则重新生成。"),
            ("内容边界", "禁止露骨色情、仇恨和过度血腥描写；暗黑风格仍保持可接受尺度。"),
        ],
        [2160, 7200],
    )

    add_heading(doc, "7.2 私人题库", level=2)
    add_bullets(
        doc,
        [
            "仅管理员能够新增、修改、停用和删除题目。",
            "每道题保存汤面、汤底和关键真相，不要求标注难度或风格。",
            "支持 JSON 批量导入和导出题库备份。",
            "房主选择私人题库时直接随机抽取，不显示筛选设置。",
            "默认避开最近抽取的 10 道题；题目不足时选择最久未出现者。",
            "普通玩家和房主均不能读取汤底，答案只由服务端发送给 AI。",
        ],
    )

    add_section_break(doc)
    add_running_furniture(doc.sections[-1])
    add_heading(doc, "8. 视觉与体验规范")
    add_callout(
        doc,
        "视觉方向",
        "深夜寝室感 + 现代极简布局。温暖、安静、清晰，不采用廉价恐怖风或复杂游戏 HUD。",
    )
    add_table(
        doc,
        ["设计层", "规范"],
        [
            ("色彩", "深蓝灰作为空间底色，暖琥珀作为行动与氛围强调色，正文保持高对比。"),
            ("排版", "仅简体中文；标题清晰，正文不过密，问题和回答形成稳定阅读节奏。"),
            ("卡片", "轻微纸张质感、柔和边界和克制阴影；不同消息类型通过层级区分。"),
            ("动效", "AI思考使用呼吸灯和短文字动画；避免大幅位移、频繁闪烁和炫技转场。"),
            ("状态反馈", "连接、同步、排队、失败和重连必须有明确文字说明，不能只依赖颜色。"),
            ("窗口", "面向 Windows 10/11 64 位桌面；支持常见笔记本分辨率与窗口缩放。"),
        ],
        [2160, 7200],
    )

    add_heading(doc, "9. 技术架构")
    add_table(
        doc,
        ["层级", "首选技术", "职责"],
        [
            ("Windows客户端", "Tauri 2 + Vue 3 + TypeScript + Vite", "窗口、界面、交互、本地历史和网络适配。"),
            ("桌面外壳", "Rust / Tauri", "安装包、系统能力、本地安全存储和应用生命周期。"),
            ("实时通信", "HTTP + WebSocket", "创建/加入、命令请求、房间事件和断线重连。"),
            ("后端", "Python + FastAPI + Uvicorn", "房间状态、串行队列、权限、AI编排和广播。"),
            ("AI", "DeepSeek API", "题目生成、主持判断、提示、结案和赛后总结。"),
            ("持久化", "客户端 SQLite；服务端 SQLite", "客户端保存本地对局；服务端只保存题库与配置。"),
            ("部署", "Ubuntu 24.04 + systemd", "通过公网 IP 提供服务，暂不使用 Docker。"),
        ],
        [1680, 3360, 4320],
    )

    add_heading(doc, "9.1 客户端分层", level=2)
    add_bullets(
        doc,
        [
            "Presentation：Vue页面、组件、主题和动效。",
            "Application：房间状态机、问题队列、结案流程和错误处理。",
            "Transport：统一通信接口，分别实现 MockTransport 与 ServerTransport。",
            "Persistence：本地身份、设置和对局历史。",
            "Protocol：共享类型、事件名称、版本号和数据校验。",
        ],
    )
    add_body(
        doc,
        "开发阶段默认使用 MockTransport，支持虚拟玩家和模拟 AI 回答，从而在没有后端时完整验证用户流程。联调阶段只切换传输实现，不重写页面。",
    )

    add_heading(doc, "9.2 服务端数据策略", level=2)
    add_table(
        doc,
        ["数据类型", "存储位置", "生命周期"],
        [
            ("活跃房间与成员", "服务端内存", "房间关闭、结算或服务器重启时清除。"),
            ("讨论、正式问答与提示", "服务端内存", "仅用于当前房间同步和 AI 上下文。"),
            ("私人题库", "服务端 SQLite", "管理员主动修改或删除。"),
            ("近期抽题记录", "服务端 SQLite", "用于避免重复，不关联昵称或玩家。"),
            ("本地对局历史", "每个客户端 SQLite", "由本机用户在程序内查看或删除。"),
        ],
        [2400, 2880, 4080],
    )

    add_heading(doc, "10. 通信协议方向")
    add_body(
        doc,
        "HTTP负责低频命令与管理操作，WebSocket负责房间事件。所有消息包含协议版本、事件 ID、房间 ID、服务端时间和结构化载荷。客户端必须能安全忽略未知字段，为后续兼容升级留出空间。",
    )
    add_table(
        doc,
        ["通道", "代表能力"],
        [
            ("HTTP", "健康检查、创建房间、加入房间、管理员验证、题库 CRUD、导入与导出。"),
            ("WebSocket 命令", "发送讨论、提交问题、撤回问题、请求提示、发起结案、提交推理、放弃和关闭房间。"),
            ("WebSocket 事件", "成员变化、队列更新、AI开始思考、AI回答、公共提示、结案结果、房主转移和房间关闭。"),
        ],
        [2160, 7200],
    )

    add_heading(doc, "11. 安全边界")
    add_callout(
        doc,
        "当前限制",
        "首版通过公网 IP 使用未加密 HTTP/WebSocket。网络链路上的第三方理论上可读取或篡改流量，因此管理员密码必须专用且不得与其他账号共用。",
        fill=PALE_AMBER,
    )
    add_bullets(
        doc,
        [
            "DeepSeek API Key只保存在服务端环境变量，不进入客户端、安装包、日志或文档示例。",
            "管理员密码在服务端只保存强哈希；登录会话设置短期失效时间。",
            "房间邀请码不是安全凭据，只用于小范围好友加入。",
            "AI返回的结构化数据必须校验枚举、长度和必填字段。",
            "协议与配置预留 HTTP→HTTPS、WS→WSS 的无破坏切换路径。",
        ],
    )

    add_heading(doc, "12. Windows客户端交付物")
    add_bullets(
        doc,
        [
            "完整 Tauri/Vue/TypeScript 源代码与关键逻辑注释。",
            "可独立运行的本地模拟模式。",
            "Windows 10/11 x64 的 Trepang Soup Setup.exe。",
            "设计令牌、组件状态和页面交互说明。",
            "本地开发、测试和打包说明。",
            "HTTP接口、WebSocket事件及共享数据模型文档。",
            "Ubuntu 24.04 后端实施、WSL联调和公网部署文档。",
            "DeepSeek主持人提示词、结构化输出和错误恢复设计。",
        ],
    )

    add_heading(doc, "13. 实施阶段")
    add_table(
        doc,
        ["阶段", "核心产出", "完成标准"],
        [
            ("1. 基础工程", "Tauri/Vue工程、路由、状态层、设计令牌。", "开发模式可启动，生产前端可构建。"),
            ("2. 房间入口", "昵称、创建、加入、等候大厅。", "模拟服务可完成创建/加入/房主操作。"),
            ("3. 游戏核心", "三栏布局、讨论、问题队列、AI回答。", "队列顺序、撤回和状态转换正确。"),
            ("4. 提示与结案", "公共提示、完整推理、放弃、结算。", "可从开局完整走到赛后总结。"),
            ("5. 本地能力", "身份恢复、对局历史、题库管理界面。", "重启客户端后可查看本地记录。"),
            ("6. 协议与文档", "后端契约、AI编排、部署指南。", "可在 WSL 中按文档独立实现后端。"),
            ("7. 发行验证", "安装包、构建记录、UI与流程回归。", "Windows x64 安装、启动和卸载正常。"),
        ],
        [1440, 3960, 3960],
    )

    add_heading(doc, "14. 首版验收标准")
    add_bullets(
        doc,
        [
            "新玩家无需注册即可在一分钟内进入房间。",
            "模拟模式能够完整演示多人加入、讨论、提问、撤回、回答、提示、结案和结算。",
            "正式问题严格按服务器顺序处理，不出现重复回答或错误撤回。",
            "中途加入者能获得完整房间快照，断线重连不会重复消息。",
            "汤底在结案前不出现在普通客户端数据中。",
            "本地历史在程序内可读，服务端不保留完成对局。",
            "界面在常见 Windows 10/11 x64 环境中清晰、流畅、无明显闪烁或布局溢出。",
            "DeepSeek Key不进入任何客户端产物。",
            "安装包可完成安装、启动、覆盖升级和卸载。",
        ],
    )

    add_heading(doc, "15. 主要风险与处理")
    add_table(
        doc,
        ["风险", "影响", "处理策略"],
        [
            ("AI判断不稳定", "回答前后矛盾或错误结案。", "结构化输出、完整上下文、二次校验和可控重试。"),
            ("AI生成题目不自洽", "玩家无法合理推出答案。", "生成后独立质量检查，失败则自动重生成。"),
            ("公网IP未加密", "内容和管理会话可能被截获。", "使用专用密码、限制管理会话，并保留TLS升级边界。"),
            ("服务器重启", "活跃房间丢失。", "明确提示当前限制；未来可选短期快照，但首版不持久化对局。"),
            ("客户端与后端漂移", "联调时事件字段不一致。", "协议版本、共享Schema、MockTransport与契约测试。"),
            ("界面信息密度过高", "小屏幕上难以阅读。", "三栏自适应、明确视觉层级、窄窗口切换分区标签。"),
        ],
        [2400, 3120, 3840],
    )

    add_heading(doc, "16. 已确认的产品决策")
    add_bullets(
        doc,
        [
            "正式名称为 Trepang Soup，中文副标题为“海龟汤”。",
            "界面仅支持简体中文，目标平台为 Windows 10/11 64 位。",
            "后端最终部署到 Ubuntu 24.04 x86-64 公网服务器。",
            "先完成 Windows 客户端与后端文档，再在 WSL Ubuntu 24.04 中实现服务端。",
            "客户端以好友房间、纯文字、合作推理和极致便捷为唯一主线。",
            "本规划作为首版实施基线；新增需求进入后续版本，不隐式扩大首版范围。",
        ],
    )

    add_callout(
        doc,
        "下一步",
        "初始化 Windows 客户端工程，先完成设计系统、首页、创建/加入房间和可运行的模拟大厅。",
        fill=PALE_BLUE,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.core_properties.title = "Trepang Soup 项目规划"
    doc.core_properties.subject = "海龟汤多人在线桌面平台产品与实施规划"
    doc.core_properties.author = "Trepang Soup Project"
    doc.core_properties.keywords = "Trepang Soup, 海龟汤, Tauri, DeepSeek, 项目规划"
    doc.save(OUTPUT)
    print("Project plan document generated.")


if __name__ == "__main__":
    main()
